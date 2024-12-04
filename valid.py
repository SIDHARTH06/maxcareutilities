from contextlib import nullcontext
import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from fuzzywuzzy import fuzz
from fuzzywuzzy import process
import os
import math
# Sidebar for selecting which tool to use
st.sidebar.title("Choose Tool")
app_selection = st.sidebar.radio("Select an option", ("Excel to Word Generator", "School Names Matching Tool"))

# Excel to Word Generator Tool
if app_selection == "Excel to Word Generator":
    # Initialize session state for tracking serial numbers
    if 'serial_numbers' not in st.session_state:
        st.session_state.serial_numbers = []

    st.title("Excel to Word Generator")

    # Address Formatting Function
    import re

    def format_address(address):
        """Format the address by:
           1. Replacing all commas with a newline.
           2. Replacing sequences of 3 or more spaces with a newline.
           3. Replacing consecutive newlines with a single newline.
           4. Adding a newline after "SCHOOL", "Academy", "Vidyalaya", or "HSS" if not already present.
        """
        # Step 1: Replace all commas with a newline
        address = address.replace(',', '\n')

        # Step 2: Replace sequences of 3 or more spaces with a newline
        address = re.sub(r' {3,}', '\n', address)  # Matches 3 or more consecutive spaces

        # Step 3: Replace consecutive newlines with a single newline
        address = re.sub(r'\n+', '\n', address)  # Replaces multiple newlines with a single newline

        # Step 4: Add a newline after "SCHOOL", "Academy", "Vidyalaya", or "HSS" (case insensitive) if no newline exists
        address = re.sub(r'(?i)(SCHOOL|Academy|Vidyalaya|HSS)(?!\s*\n)', r'\1\n', address)

        return address.strip()



    # File upload
    uploaded_file = st.file_uploader("Upload an Excel File", type=['xlsx'])

    if uploaded_file:
        # Load the Excel file into a DataFrame
        df = pd.read_excel(uploaded_file)
        st.write("Uploaded Data Preview:")
        st.dataframe(df)

        # Range input for serial numbers
        serial_range = st.text_input("Enter Serial Numbers Range (e.g., 1-5):")

        # Button to add range
        if st.button("Add Range"):
            try:
                start, end = map(int, serial_range.split('-'))
                if start > 0 and end >= start:
                    st.session_state.serial_numbers.extend(range(start, end + 1))
                    st.success(f"Added Serial Numbers: {list(range(start, end + 1))}")
                else:
                    st.error("Invalid range! Ensure the range is positive and properly formatted.")
            except ValueError:
                st.error("Invalid format! Use the format 'start-end' (e.g., 1-5).")

        # Display added serial numbers
        if st.session_state.serial_numbers:
            st.write("Added Serial Numbers:", sorted(set(st.session_state.serial_numbers)))

        # Button to generate Word file
        if st.button("Generate Word File"):
            # Filter DataFrame for selected serial numbers
            selected_rows = df[df['SL'].isin(st.session_state.serial_numbers)]
            if selected_rows.empty:
                st.error("No matching data for the selected serial numbers!")
            else:
                document = Document()
                # Ensure all data in the 'ADDRESS' column is treated as strings
                selected_rows['ADDRESS'] = selected_rows['ADDRESS'].astype(str).apply(format_address)

                for address in selected_rows['ADDRESS']:
                    # Check if address is not empty, NaN, or null before processing
                    if address.strip() == "" or pd.isna(address) or address == 'nan':  # Skip empty or NaN addresses
                        continue

                    # Create a table with one cell for the address box
                    table = document.add_table(rows=1, cols=1)
                    table.autofit = False  # Disable auto resizing for custom styling

                    # Set the width of the table cell
                    cell = table.cell(0, 0)
                    cell.text = address  # Add formatted address

                    # Apply font formatting
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Black font color
                            run.font.size = Pt(12)  # Font size

                    # Add a black border with 1pt thickness to the table cell
                    tc_pr = cell._element.get_or_add_tcPr()
                    if tc_pr is None:
                        tc_pr = OxmlElement("w:tcPr")  # Create tcPr element if missing
                        cell._element.append(tc_pr)

                    # Set the border properties with 1pt thickness and black color
                    for border_name in ["top", "bottom", "left", "right"]:
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn("w:val"), "single")  # Solid border
                        border.set(qn("w:sz"), "1")  # Border thickness set to 1pt
                        border.set(qn("w:space"), "0")
                        border.set(qn("w:color"), "000000")  # Black color
                        tc_pr.append(border)

                    # Enable word wrapping in the table cell
                    no_wrap = OxmlElement("w:noWrap")
                    no_wrap.set(qn("w:val"), "false")
                    tc_pr.append(no_wrap)

                    # document.add_paragraph()  # Add space between boxes

                # Save the Word document to a BytesIO object
                buffer = BytesIO()
                document.save(buffer)
                buffer.seek(0)

                # Provide download option
                st.download_button(
                    label="Download Word File",
                    data=buffer,
                    file_name="generated_addresses.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# School Names Matching Tool
if app_selection == "School Names Matching Tool":
    # Function to perform fuzzy matching
    def is_match(school_name, check_schools, threshold):
        if not isinstance(school_name,float):
            matches = process.extract(school_name, check_schools, scorer=fuzz.ratio)
            return any(match[1] >= threshold for match in matches)

    st.title('Maxcare India School Names Matching Tool')

    # Upload input Excel sheet and check sheet
    input_file = st.file_uploader("Upload Input Sheet (Excel)", type=['xlsx'])
    check_file = st.file_uploader("Upload Check Sheet (Excel)", type=['xlsx'])

    # If the files are uploaded
    if input_file is not None and check_file is not None:
        # Read the input and check sheets
        input_df = pd.read_excel(input_file)
        check_df = pd.read_excel(check_file)

        # Ensure the column names are correctly recognized (case sensitive check)
        if 'Name' not in input_df.columns or 'INSTITUTE' not in check_df.columns:
            st.error("The input sheet must have a 'Name' column and the check sheet must have an 'INSTITUTE' column.")
        else:
            # Slider to set the threshold for fuzzy matching
            threshold = st.slider("Select Fuzzy Matching Threshold (%)", 0, 100, 85)

            # Extract school names from both sheets
            input_schools = input_df['Name'].tolist()
            check_schools = check_df['INSTITUTE'].tolist()

            # Find unmatched schools
            input_schools = [x for x in input_schools if not (isinstance(x, float) and math.isnan(x) or x == "nan")]
            check_schools = [x for x in check_schools if not (isinstance(x, float) and math.isnan(x) or x == "nan")]
            unmatched_rows = input_df[~input_df['Name'].apply(lambda x: is_match(x, check_schools, threshold))]

            # Save the output file with the same name but with '_processed' suffix
            output_filename = os.path.splitext(input_file.name)[0] + '_processed.xlsx'
            unmatched_rows.to_excel(output_filename, index=False)

            # Provide a download link for the processed file
            st.write(f"Processed file saved as: {output_filename}")
            st.download_button(
                label="Download Processed File",
                data=open(output_filename, 'rb').read(),
                file_name=output_filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
